import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python docx not installed.")
    sys.exit(1)

doc = Document()

# Cover Page
for _ in range(4):
    doc.add_paragraph()

title = doc.add_heading('INJECTION VULNERABILITIES', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('An Analytical & Real World Perspective')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n')

unit_code = doc.add_paragraph()
unit_code.add_run('UNIT CODE: ').bold = True
unit_code.add_run('BIT 425E')
unit_code.alignment = WD_ALIGN_PARAGRAPH.CENTER

unit_name = doc.add_paragraph()
unit_name.add_run('UNIT NAME: ').bold = True
unit_name.add_run('WEB APPLICATION SECURITY')
unit_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n')

g = doc.add_paragraph()
g.add_run('GROUP MEMBERS').bold = True
g.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Kept the student ID hyphens intact so the IDs format correctly.
members = [
    "Brian Ndinya       SIT/B/01-02627/2022",
    "Bugendi Devis      SIT/B/01-02891/2022",
    "Elshalom Baraka    SIT/B/01-02926/2022",
    "Benard Rorat       SIT/B/01-02893/2022",
    "Emmanuel Muchina   SIT/B/01-02896/2022"
]

for m_text in members:
    p = doc.add_paragraph(m_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()


# Section 1
doc.add_heading('1. Conceptual Explanation', level=1)
doc.add_paragraph('At its core, an injection flaw is a critical problem of mistaken identity between data and instructions. When you enter information into a web application, such as submitting your username or searching for a product, the system is supposed to treat that input purely as data. However, if the application is not designed securely, it might accidentally read your raw input as a direct system command.')
doc.add_paragraph('For example, an attacker can submit a cleverly crafted piece of malicious code instead of a normal search word. Because the system incorrectly fails to separate the incoming user data from its own internal logic, it blindly executes the attacker code as if it were a legitimate instruction. This flaw is incredibly dangerous because it essentially allows an external attacker to reprogram the application on the fly, enabling them to bypass login screens, access highly sensitive databases, or manipulate the underlying core operating systems without permission.')

# Section 2
doc.add_heading('2. Causes', level=1)
doc.add_paragraph('The primary causes of injection vulnerabilities involve practices where user input is implicitly trusted or not correctly sanitized. Principal technical causes include:')

p = doc.add_paragraph()
p.add_run('String Concatenation over Parameterization: ').bold = True
p.add_run('This happens when developers build database queries by directly combining lines of code with raw user input. Instead of securely isolating the input, the application merges everything together, creating a direct pathway for malicious commands.')

p = doc.add_paragraph()
p.add_run('Insufficient Input Validation: ').bold = True
p.add_run('This occurs when an application accepts user provided data without strictly verifying its type, length, or format against an expected list of safe values.')

p = doc.add_paragraph()
p.add_run('Flawed Implementation in ORMs: ').bold = True
p.add_run('Even when developers employ modern Object Relational Mappers to handle databases securely, they sometimes bypass these safe tools to write direct, high risk database queries purely out of convenience.')

p = doc.add_paragraph()
p.add_run('Incorrect Privilege Configuration: ').bold = True
p.add_run('This involves operating applications with far more permissions than they actually need, such as logging into a database with full administrative rights. While excessive privileges do not create the vulnerability itself, they drastically multiply the ultimate damage an attacker can do once inside.')


# Section 3
doc.add_heading('3. Attack Scenarios', level=1)
doc.add_paragraph('Attackers exploit injection vulnerabilities through various technical methods adapted to the specific application logic:')


p = doc.add_paragraph()
p.add_run('Authentication Bypass: ').bold = True
p.add_run("During the login process, an attacker inputs a specific logic condition into a typical username field. The resulting backend verification evaluates to universally true, allowing the attacker to completely bypass the password check and immediately gain access to the system.")

p = doc.add_paragraph()
p.add_run('Data Exfiltration via Inference (Blind SQLi): ').bold = True
p.add_run('When applications hide basic error messages, attackers resort to blind injection. They inject scripts that ask the database true or false questions and measure exactly how long the application takes to respond. This allows them to slowly and silently extract sensitive strings character by character.')

p = doc.add_paragraph()
p.add_run('OS Command Execution: ').bold = True
p.add_run("If an application interacts with its underlying operating system, taking actions like processing an uploaded file, an attacker can append system commands to their input. The host system then executes both the legitimate and the malicious commands, potentially exposing critical internal files.")


# Section 4
doc.add_heading('4. Impact Analysis', level=1)
doc.add_paragraph('The consequences of a successful injection attack typically extend across several critical business operations:')


p = doc.add_paragraph()
p.add_run('Data Breaches: ').bold = True
p.add_run('Loss of confidentiality is immediate. Threat actors can extract entire databases containing proprietary intellectual property, employee details, or sensitive customer data such as financial or health records.')

p = doc.add_paragraph()
p.add_run('Data Integrity and Service Disruption: ').bold = True
p.add_run('Beyond simple theft, attackers can modify financial transaction ledgers, delete essential records, or execute destructive commands that permanently corrupt organizational data.')

p = doc.add_paragraph()
p.add_run('Regulatory and Reputational Consequences: ').bold = True
p.add_run('A confirmed breach usually triggers severe consequences, resulting in heavy regulatory fines, costly digital forensics, and a prolonged loss of consumer and partner trust.')


# Section 5
doc.add_heading('5. Countermeasures and Mitigations', level=1)
doc.add_paragraph('Preventing injection requires a systematic, secure design approach rather than relying solely on perimeter defenses. Reliable countermeasures include:')

p = doc.add_paragraph()
p.add_run('Parameterized Queries (Prepared Statements): ').bold = True
p.add_run('This represents the most effective primary defense. By strictly separating the query logic from the user supplied data, the backend safely processes the input purely as a literal value, making it impossible to execute as code.')

p = doc.add_paragraph()
p.add_run('Strict Allowlisting for Input Validation: ').bold = True
p.add_run('Rather than trying to anticipate and block malicious inputs, applications should enforce strict definitions of acceptable data formats, immediately rejecting any request that does not perfectly conform.')

p = doc.add_paragraph()
p.add_run('The Principle of Least Privilege: ').bold = True
p.add_run('Organizations must ensure web applications connect to databases using accounts with only the bare minimum permissions required. A read only service account successfully prevents an attacker from altering internal structures or dropping tables even if an injection flaw exists.')


# Section 6
doc.add_heading('6. Real World Case Studies and Fact Checks', level=1)
doc.add_paragraph('To contextualize this vulnerability analytically, we must examine documented incidents and formally verify their narratives:')

p = doc.add_paragraph()
p.add_run('✅ The Kenyan Ministry of Foreign Affairs (2016) Partially True').bold = True

doc.add_paragraph('The core claim that hackers affiliated with Anonymous breached the Kenyan Ministry of Foreign Affairs in April 2016 and claimed to have stolen 1 Terabyte of data is true. A sample of 95 documents was indeed released online, including routine diplomatic communications and security discussions.')

doc.add_paragraph('However, two key details are disputed:')

p = doc.add_paragraph()
p.add_run('The Method: ').bold = True
p.add_run('While cybersecurity discussions often attribute this to a classic SQL Injection, the Kenyan government ICT Cabinet Secretary stated the attack was actually a phishing campaign. Officials reported that employees were tricked into clicking malicious links giving hackers access to email credentials, rather than a direct technical breach of the database server.')

p = doc.add_paragraph()
p.add_run('The Impact: ').bold = True
p.add_run('Although the exposure of critical vulnerabilities was claimed by hackers, the government downplayed the severity, stating that no classified or Top Secret material was accessed and that the documents were mostly routine correspondence.')

p = doc.add_paragraph()
p.add_run('Verdict: ').bold = True
p.add_run('The event happened, but the specific method (SQLi versus Phishing) is actively disputed by official sources.')

p = doc.add_paragraph()
p.add_run('✅ The Silent Threat to Kenyan SACCOs True in Principle').bold = True

doc.add_paragraph('This accurately reflects a well documented and serious vulnerability in the Kenya financial sector. While finding a specific, globally publicized case of a hacker manipulating SQL to inflate balances and withdraw via M PESA is rare due to non disclosure, the systemic risk is confirmed by official sources.')

doc.add_paragraph('The Sacco Societies Regulatory Authority (SASRA) has issued multiple warnings confirming that 98 percent of cyber threats to SACCOs come through third party vendors, such as those providing mobile money API integrations. SASRA has specifically warned that digital credit channels and pay bill accounts are at extremely high risk. Furthermore, the regulator has ordered SACCOs to physically secure their systems against the specific scenario of manipulating member portals to exploit financial platforms.')

p = doc.add_paragraph()
p.add_run('Verdict: ').bold = True
p.add_run('The scenario is a highly recognized and active threat in Kenya, heavily supported by official regulatory warnings.')

p = doc.add_paragraph()
p.add_run('✅ Global Benchmark The MOVEit Zero Day (2023) True').bold = True

doc.add_paragraph('This statement is completely accurate. The 2023 MOVEit Transfer breach is a textbook example of the devastation caused by an authentic SQL injection vulnerability.')

p = doc.add_paragraph()
p.add_run('The Vulnerability: ').bold = True
p.add_run('The attack exploited a zero day SQL injection vulnerability officially cataloged as CVE 2023 34362.')

p = doc.add_paragraph()
p.add_run('The Attacker: ').bold = True
p.add_run('The CL0P ransomware group was solely responsible for the campaign.')

p = doc.add_paragraph()
p.add_run('The Impact: ').bold = True
p.add_run('The attack had a massive supply chain effect. Official reports confirm it affected over 2,600 companies and more than 83 million individuals, making it one of the largest data breaches of the year.')

p = doc.add_paragraph()
p.add_run('Verdict: ').bold = True
p.add_run('This is an incredibly accurate description of a major global cybersecurity event.')

p = doc.add_paragraph()
p.add_run('💡 Key Takeaways').bold = True

doc.add_paragraph('SQL Injection is Real: The MOVEit case proves that even expensive, enterprise software is deeply vulnerable to SQL injection if security is not an absolute priority.')

doc.add_paragraph('Fintech Risk is Confirmed: Kenyan regulators are actively warning about the risks of integrating mobile money with external systems, validating the factual reality of the Silent Threat.')

doc.add_paragraph('Verify the Method: The summary of the 2016 Kenya hack is an excellent example of how a cyberattack narrative can be reported as an SQL injection when the underlying cause may actually be human error (phishing). Both lead to devastating data theft, but the method dictates the security defense required.')


# Section 7
doc.add_heading('7. References', level=1)
doc.add_paragraph('1. Open Worldwide Application Security Project (OWASP). A03:2021 Injection. OWASP Top 10.')
doc.add_paragraph('2. Sacco Societies Regulatory Authority (SASRA) Kenya. Annual Sacco Supervision Reports and Cyber Security Guidelines for SACCOs.')
doc.add_paragraph('3. National Institute of Standards and Technology (NIST). CVE 2023 34362 Detail. National Vulnerability Database. (Pertaining to the MOVEit Transfer Zero Day vulnerability).')
doc.add_paragraph('4. Public advisories and Kenya Ministry of ICT statements regarding the April 2016 unauthorized access incident at the Ministry of Foreign Affairs.')


out_path = '/home/wakanda_forever/Desktop/AISHA/documentations/injection.docx'
os.makedirs(os.path.dirname(out_path), exist_ok=True)
doc.save(out_path)
print(f"Successfully saved {out_path}")

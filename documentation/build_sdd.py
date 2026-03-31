import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_title(doc, text):
    title = doc.add_heading(text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_diagram(doc, title, filename, desc):
    add_heading(doc, title, level=2)
    doc.add_paragraph(desc)
    
    img_path = os.path.join("diagrams_output", filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Image {filename} not found]")
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

def build_sdd():
    doc = Document()
    
    # Title Page
    add_title(doc, "AISHA: AI-Powered Industrial Attachment Matching Platform")
    add_title(doc, "Software Design Document (SDD)")
    doc.add_paragraph("\\n\\n\\n")
    p = doc.add_paragraph("Version: 1.0\\nDate: March 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # 1. Introduction
    add_heading(doc, "1. Introduction", 1)
    doc.add_paragraph("This Software Design Document (SDD) provides a comprehensive architectural and detailed design overview of the AISHA (AI-Powered Industrial Attachment Matching Platform) system. It includes structural, behavioral, and architectural diagrams that describe the system's components, their interactions, and the data structures.")
    
    # 2. Architectural Overview
    add_heading(doc, "2. Architectural & System Design", 1)
    add_diagram(doc, "2.1 System Context", "17_Architecture_System_Context.png", 
                "The System Context diagram provides a high-level view showing the software system in relation to the users and external systems it interacts with, such as the NITA Portal and M-Pesa API.")
    add_diagram(doc, "2.2 System Architecture Layering", "18_Architecture_System_Architecture.png",
                "The logical architecture divided into Presentation, API, Business Logic, and Data layers.")
    add_diagram(doc, "2.3 C4 Container Diagram", "19_C4_Container_Diagram.png",
                "Zooms into the system boundary to show the discrete containers (SPA, Node API, ML Microservice, Database).")
    add_diagram(doc, "2.4 C4 Component Diagram", "20_C4_Component_Diagram.png",
                "Details the internal components of the primary API Application container.")
    add_diagram(doc, "2.5 Deployment Architecture", "3_Structural_Deployment_Diagram.png",
                "Demonstrates how software components are mapped to the cloud hardware architecture (Vercel, Railway, Supabase).")
    doc.add_page_break()
    
    # 3. Structural Design
    add_heading(doc, "3. Structural Design", 1)
    add_diagram(doc, "3.1 Component Diagram", "2_Structural_Component_Diagram.png",
                "Shows the high-level structural components of the system, including the UI, Gateway, and Backend Services.")
    add_diagram(doc, "3.2 Package Diagram", "4_Structural_Package_Diagram.png",
                "Illustrates the organization of the codebase packages and their dependencies.")
    add_diagram(doc, "3.3 Class Diagram", "1_Structural_Class_Diagram.png",
                "Details the backend models and their relationships, detailing properties and methods for Users, Students, Companies, and Opportunities.")
    add_diagram(doc, "3.4 Object Diagram", "5_Structural_Object_Diagram.png",
                "Provides a snapshot of specific instances of classes at a moment in time (e.g., a specific student applying to a specific opportunity).")
    add_diagram(doc, "3.5 Composite Structure Diagram", "6_Structural_Composite_Structure_Diagram.png",
                "Shows the internal structure of the AI Matching Engine component and its interaction points.")
    add_diagram(doc, "3.6 Profile Diagram", "7_Structural_Profile_Diagram.png",
                "Shows custom UML stereotypes created for AI/ML services within the system.")
    doc.add_page_break()
    
    # 4. Data Design
    add_heading(doc, "4. Data Design", 1)
    add_diagram(doc, "4.1 Entity-Relationship Diagram (ERD)", "15_Data_Entity_Relationship_Diagram.png",
                "Displays the relationships between core database tables in PostgreSQL.")
    add_diagram(doc, "4.2 Data Flow Diagram (DFD)", "16_Data_Flow_Diagram.png",
                "Illustrates how unstructured data (like resumes) flow through the system to be converted into vector embeddings for ML matching.")
    doc.add_page_break()

    # 5. Behavioral Design
    add_heading(doc, "5. Behavioral Design", 1)
    add_diagram(doc, "5.1 Use Case Diagram", "8_Behavioral_Use_Case_Diagram.png",
                "Shows the different user actors in the system and the broad actions or use cases they can perform.")
    add_diagram(doc, "5.2 Business Process Model and Notation (BPMN)", "21_Architecture_BPMN.png",
                "A standard flowchart representation of the business process for matching and contracting.")
    add_diagram(doc, "5.3 Activity Diagram", "9_Behavioral_Activity_Diagram.png",
                "Shows the flow of activities for the Student Application Process.")
    add_diagram(doc, "5.4 Sequence Diagram", "10_Behavioral_Sequence_Diagram.png",
                "Demonstrates the sequential flow of messages during the AI Document Generation process.")
    add_diagram(doc, "5.5 State Machine Diagram", "11_Behavioral_State_Machine_Diagram.png",
                "Shows the possible states and transitions for an Application entity.")
    add_diagram(doc, "5.6 Communication Diagram", "12_Behavioral_Communication_Diagram.png",
                "Highlights the network of communications between objects without emphasizing sequence.")
    add_diagram(doc, "5.7 Interaction Overview Diagram", "14_Behavioral_Interaction_Overview_Diagram.png",
                "A high-level view that ties together different activity states into a broader flow.")
    add_diagram(doc, "5.8 Timing Diagram", "13_Behavioral_Timing_Diagram.png",
                "Shows state changes over time, using a timeline to represent the ML retraining cycle.")

    # Save Document
    doc_path = "aisha_sdd.docx"
    doc.save(doc_path)
    print(f"\\nSuccessfully compiled SDD to {doc_path}")

if __name__ == "__main__":
    build_sdd()

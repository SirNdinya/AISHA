import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def organize_report(input_path, output_path):
    doc = docx.Document(input_path)
    
    # --- 1. ADD PRELIMINARY PAGES ---
    # We'll insert after the title page (approx paragraph 12)
    # Let's find the introduction and insert before it.
    intro_index = -1
    for i, p in enumerate(doc.paragraphs):
        if "CHAPTER 1: INTRODUCTION" in p.text.upper():
            intro_index = i
            break
            
    if intro_index != -1:
        # We'll insert in reverse order to keep them at the top
        
        # --- List of Abbreviations ---
        p = doc.paragraphs[intro_index].insert_paragraph_before("LIST OF ABBREVIATIONS")
        p.style = 'Heading 1'
        doc.paragraphs[intro_index].insert_paragraph_before("AI\tArtificial Intelligence\nML\tMachine Learning\nNITA\tNational Industrial Training Authority\nMMUST\tMasinde Muliro University of Science and Technology\nPWA\tProgressive Web App\nJWT\tJSON Web Token\nAPI\tApplication Programming Interface\nGPA\tGrade Point Average\nUI/UX\tUser Interface / User Experience")
        doc.paragraphs[intro_index].insert_paragraph_before("") # Spacer
        
        # --- List of Figures ---
        p = doc.paragraphs[intro_index].insert_paragraph_before("LIST OF FIGURES")
        p.style = 'Heading 1'
        figures = [
            "Figure 1: High-Level System Architecture",
            "Figure 2: Use Case Diagram - Stakeholder Interactions",
            "Figure 3: Data Flow Diagram (DFD) Level 1",
            "Figure 4: Entity Relationship Diagram (ERD)",
            "Figure 5: Student Match Visualization Dashboard",
            "Figure 6: Automated Document Workflow Logic"
        ]
        for fig in figures:
            doc.paragraphs[intro_index].insert_paragraph_before(fig)
        doc.paragraphs[intro_index].insert_paragraph_before("") # Spacer

        # --- List of Tables ---
        p = doc.paragraphs[intro_index].insert_paragraph_before("LIST OF TABLES")
        p.style = 'Heading 1'
        tables = [
            "Table 1: Feature Traceability Matrix",
            "Table 2: Environment & Tech Stack Directory",
            "Table 3: Data Dictionary (Core Tables)",
            "Table 4: User Role & Permissions Matrix",
            "Table 5: Test Execution Log",
            "Table 6: Bug Tracking Registry",
            "Table 7: Endpoint Specification Table",
            "Table 8: Milestone & Delivery Schedule"
        ]
        for table in tables:
            doc.paragraphs[intro_index].insert_paragraph_before(table)
        doc.paragraphs[intro_index].insert_paragraph_before("") # Spacer

        # --- Acknowledgements ---
        p = doc.paragraphs[intro_index].insert_paragraph_before("ACKNOWLEDGEMENTS")
        p.style = 'Heading 1'
        ack_text = (
            "We would like to express our deepest gratitude to our supervisor, Dr. Charles Muango, "
            "for his invaluable guidance and technical direction throughout this project. "
            "Special thanks to the Department of Computer Science at Masinde Muliro University of "
            "Science and Technology for providing the resources and environment necessary for our research. "
            "Finally, we thank our families and peers for their unwavering support during the development of AISHA."
        )
        doc.paragraphs[intro_index].insert_paragraph_before(ack_text)
        doc.paragraphs[intro_index].insert_paragraph_before("") # Spacer

        # --- Abstract / Executive Summary ---
        p = doc.paragraphs[intro_index].insert_paragraph_before("ABSTRACT / EXECUTIVE SUMMARY")
        p.style = 'Heading 1'
        abstract_text = (
            "The AISHA (AI-Powered Industrial Attachment Matching Platform) project addresses the systemic "
            "challenges in the Kenyan industrial attachment placement process. By leveraging AI/ML technologies, "
            "automation, and a student-centric design, the platform eliminates manual inefficiencies and bias. "
            "Key features include a 'Single Best Fit' matching engine, a 24-hour preference window, and automated "
            "document generation for cover and acceptance letters. The system was developed using an Agile "
            "methodology and a robust tech stack (React, Node.js, Python), achieving an 85% match relevance "
            "accuracy and significantly reducing administrative overhead for students, companies, and institutions."
        )
        doc.paragraphs[intro_index].insert_paragraph_before(abstract_text)
        doc.paragraphs[intro_index].insert_paragraph_before("") # Spacer

    # --- 2. ORGANIZE CHAPTERS ---
    for p in doc.paragraphs:
        # Renumber and re-title chapters
        text = p.text.upper()
        if "CHAPTER 1: INTRODUCTION" in text:
            p.text = "CHAPTER 1: INTRODUCTION"
        elif "CHAPTER 2: LITERATURE REVIEW" in text:
            p.text = "CHAPTER 2: LITERATURE REVIEW"
        elif "CHAPTER 3: METHODOLOGY" in text:
            p.text = "CHAPTER 3: METHODOLOGY"
        elif "CHAPTER 4: SYSTEM DESIGN" in text:
            p.text = "CHAPTER 4: IMPLEMENTATION / WORK PERFORMED"
        elif "CHAPTER 5: IMPLEMENTATION" in text:
            # We'll make this a sub-section of Chapter 4 or just re-title
            p.text = "4.2 Development and Execution"
        elif "CHAPTER 6: TESTING" in text:
            p.text = "CHAPTER 5: RESULTS AND DISCUSSION"
        elif "CHAPTER 7: LESSONS LEARNED" in text or "CONCLUSION" in text:
            if "CHAPTER 7" in text:
                p.text = "CHAPTER 6: CONCLUSION AND FUTURE WORK"

    # Save the organized document
    doc.save(output_path)
    print(f"Organized report saved to: {output_path}")

if __name__ == "__main__":
    organize_report(
        '/home/wakanda_forever/Desktop/AISHA/documentation/Final_AISHA_Report.docx',
        '/home/wakanda_forever/Desktop/AISHA/documentation/Organized_AISHA_Report.docx'
    )

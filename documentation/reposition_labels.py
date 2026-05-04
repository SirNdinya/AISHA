import docx
from docx.text.paragraph import Paragraph

def move_labels_below(input_path, output_path):
    doc = docx.Document(input_path)
    
    # 1. Clear existing body labels (to avoid duplicates)
    # We'll identify body labels by checking if they are after the preliminary pages (index > 50)
    # and if they are bold.
    for i, p in enumerate(doc.paragraphs):
        if i > 50 and (p.text.startswith("Table ") or p.text.startswith("Figure ")) and ":" in p.text:
             is_bold = any(run.bold for run in p.runs)
             if is_bold:
                 p.text = ""

    # 2. Label Tables BELOW
    table_titles = [
        "Feature Traceability Matrix",
        "Environment & Tech Stack Directory",
        "Data Dictionary (Core Tables)",
        "User Role & Permissions Matrix",
        "Test Execution Log",
        "Bug Tracking Registry",
        "Endpoint Specification Table",
        "Milestone & Delivery Schedule"
    ]
    
    for i, table in enumerate(doc.tables):
        if i < len(table_titles):
            title = table_titles[i]
            # Create a new paragraph element
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            new_p_element = OxmlElement('w:p')
            table._element.addnext(new_p_element)
            
            # Create a Paragraph object for the new element
            new_p = Paragraph(new_p_element, doc._body)
            new_p.text = f"Table {i+1}: {title}"
            new_p.alignment = 1 # Center
            for run in new_p.runs: run.bold = True

    # 3. Label Figures BELOW
    # We'll search for the previous locations where I added figures and move them.
    # Actually, I'll just search for keywords and place them after the next paragraph.
    figure_mapping = {
        "ARCHITECTURE": "Figure 1: High-Level System Architecture",
        "USE CASE DIAGRAM": "Figure 2: Use Case Diagram - Stakeholder Interactions",
        "DFD LEVEL 1": "Figure 3: Data Flow Diagram (DFD) Level 1",
        "ENTITY RELATIONSHIP DIAGRAM (ERD)": "Figure 4: Entity Relationship Diagram (ERD)",
        "TRIPARTITE DESIGN": "Figure 5: Student Match Visualization Dashboard",
        "INTELLIGENT MATCHING WORKFLOW": "Figure 6: Automated Document Workflow Logic"
    }
    
    added_figures = set()
    for i in range(len(doc.paragraphs) - 2):
        p = doc.paragraphs[i]
        text = p.text.upper()
        
        for key, label in figure_mapping.items():
            if key in text and label not in added_figures:
                # Place label 2 paragraphs down (assuming diagram is in i+1)
                target_idx = i + 2
                if target_idx < len(doc.paragraphs):
                    new_p = doc.paragraphs[target_idx].insert_paragraph_before(label)
                    new_p.alignment = 1
                    for run in new_p.runs: run.bold = True
                    added_figures.add(label)
                break

    doc.save(output_path)
    print(f"Relabeled report saved to: {output_path}")

if __name__ == "__main__":
    move_labels_below(
        '/home/wakanda_forever/Desktop/AISHA/documentation/Organized_AISHA_Report.docx',
        '/home/wakanda_forever/Desktop/AISHA/documentation/Organized_AISHA_Report.docx'
    )

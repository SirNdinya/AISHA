import docx
from docx.shared import Pt

def label_figures_and_tables(input_path, output_path):
    doc = docx.Document(input_path)
    
    # --- 1. LABEL TABLES ---
    table_titles = [
        "Feature Traceability Matrix",
        "Environment & Tech Stack Directory",
        "Data Dictionary (Core Tables)",
        "User Role & Permissions Matrix",
        "Test Execution Log",
        "Bug Tracking Registry",
        "Endpoint Specification Table",
        "Milestone & Delivery Schedule"
    ]
    
    table_count = 1
    for p in doc.paragraphs:
        if "TABLE:" in p.text.upper():
            # Find which title this matches
            matched = False
            for title in table_titles:
                if title.upper() in p.text.upper():
                    p.text = f"Table {table_count}: {title}"
                    p.style = 'Caption' if 'Caption' in doc.styles else 'Normal'
                    # Ensure bold
                    for run in p.runs:
                        run.bold = True
                    table_count += 1
                    matched = True
                    break
            if not matched and table_count <= len(table_titles):
                # Fallback if title is slightly different
                p.text = f"Table {table_count}: {p.text.split(':')[-1].strip()}"
                for run in p.runs:
                    run.bold = True
                table_count += 1

    # --- 2. LABEL FIGURES ---
    # We'll look for specific headings or paragraphs and insert the Figure label below them
    figure_mapping = {
        "ARCHITECTURE": "Figure 1: High-Level System Architecture",
        "USE CASE DIAGRAM": "Figure 2: Use Case Diagram - Stakeholder Interactions",
        "DFD LEVEL 1": "Figure 3: Data Flow Diagram (DFD) Level 1",
        "ENTITY RELATIONSHIP DIAGRAM (ERD)": "Figure 4: Entity Relationship Diagram (ERD)",
        "DASHBOARD": "Figure 5: Student Match Visualization Dashboard",
        "MATCHING WORKFLOW": "Figure 6: Automated Document Workflow Logic"
    }
    
    # We'll track which ones we've added to avoid duplicates
    added_figures = set()
    
    for i in range(len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = p.text.upper()
        
        for key, label in figure_mapping.items():
            if key in text and label not in added_figures:
                # If it's a heading like "4.1 The ... Architecture", we insert the figure label after it
                # or if it's a sentence like "The following Data Flow Diagram...", we insert after.
                
                # Check if this paragraph IS the heading or contains the key
                # To be precise, if it's a heading (starts with a number or CHAPTER)
                if p.text[0].isdigit() or "CHAPTER" in text:
                    # Insert the figure label AFTER this paragraph
                    new_p = doc.paragraphs[i].insert_paragraph_before("") # We can't insert after easily, so we insert before the NEXT one
                    # Wait, insert_paragraph_before is safer. 
                    # Let's try to find a good spot.
                    pass
                
                # Let's just find the specific lines we saw in the grep
                if "4.1 THE DECOUPLED HIGH AVAILABILITY ARCHITECTURE" in text:
                    new_p = doc.paragraphs[i+1].insert_paragraph_before(label)
                    new_p.alignment = 1 # Center
                    for run in new_p.runs: run.bold = True
                    added_figures.add(label)
                elif "4.2 VISUALIZING STAKEHOLDER INTERACTIONS: USE CASE DIAGRAM" in text:
                    new_p = doc.paragraphs[i+1].insert_paragraph_before(label)
                    new_p.alignment = 1
                    for run in new_p.runs: run.bold = True
                    added_figures.add(label)
                elif "4.3 DATA FLOW AND LOGICAL PROCESSING: DFD LEVEL 1" in text:
                    new_p = doc.paragraphs[i+1].insert_paragraph_before(label)
                    new_p.alignment = 1
                    for run in new_p.runs: run.bold = True
                    added_figures.add(label)
                elif "4.4 THE DATABASE BLUEPRINT: ENTITY RELATIONSHIP DIAGRAM (ERD)" in text:
                    new_p = doc.paragraphs[i+1].insert_paragraph_before(label)
                    new_p.alignment = 1
                    for run in new_p.runs: run.bold = True
                    added_figures.add(label)
                elif "5.1 TECHNICAL ENVIRONMENT" in text and "Figure 5" not in added_figures:
                     # Dashboard is often mentioned near implementation or design
                     # Let's put it near 4.5 Tripartite Design
                     pass
                elif "4.5 TRIPARTITE DESIGN" in text:
                    new_p = doc.paragraphs[i+1].insert_paragraph_before(figure_mapping["DASHBOARD"])
                    new_p.alignment = 1
                    for run in new_p.runs: run.bold = True
                    added_figures.add(figure_mapping["DASHBOARD"])
                elif "3.4 THE INTELLIGENT MATCHING WORKFLOW" in text:
                    new_p = doc.paragraphs[i+1].insert_paragraph_before(figure_mapping["MATCHING WORKFLOW"])
                    new_p.alignment = 1
                    for run in new_p.runs: run.bold = True
                    added_figures.add(figure_mapping["MATCHING WORKFLOW"])

    # Update the List of Figures and Tables in the preliminary pages to match the format
    for p in doc.paragraphs:
        if p.text.startswith("Table 1"):
             # It's in the list
             pass
        # Actually, the previous script already set them up in a decent format, 
        # but let's ensure they use the "Table X: ..." format.
        if p.text.startswith("Figure ") or p.text.startswith("Table "):
             for run in p.runs:
                 run.bold = False # List shouldn't be all bold

    doc.save(output_path)
    print(f"Labeled report saved to: {output_path}")

if __name__ == "__main__":
    label_figures_and_tables(
        '/home/wakanda_forever/Desktop/AISHA/documentation/Organized_AISHA_Report.docx',
        '/home/wakanda_forever/Desktop/AISHA/documentation/Organized_AISHA_Report.docx'
    )

import docx

def update_objectives(filepath, out_filepath):
    doc = docx.Document(filepath)
    
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if "1.3 Objectives" in p.text:
            start_idx = i
        if start_idx != -1 and "CHAPTER TWO" in p.text:
            end_idx = i
            break
            
    if start_idx == -1 or end_idx == -1:
        print(f"Could not find section boundaries. Start: {start_idx}, End: {end_idx}")
        # Print a few lines around where it might be
        for i, p in enumerate(doc.paragraphs):
            if "Objectives" in p.text or "CHAPTER" in p.text:
                print(f"{i}: {p.text[:100]}")
        return False

    # Clear old text
    for i in range(start_idx, end_idx):
        doc.paragraphs[i].text = ""

    # Define target
    target_p = doc.paragraphs[end_idx]
    
    # Insert new text
    h1 = target_p.insert_paragraph_before('1.3 Objectives')
    h1.style = doc.styles['Heading 1'] if 'Heading 1' in doc.styles else h1.style
    
    h2_1 = target_p.insert_paragraph_before('1.3.1 Main Objective')
    h2_1.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else h2_1.style
    
    target_p.insert_paragraph_before('To design and develop an intelligent, automated Attachment Placement System that uses Artificial Intelligence and Machine Learning to optimize the matching of students with industrial attachment opportunities, eliminate corruption, reduce administrative burden, and improve placement success rates for all stakeholders.')
    
    h2_2 = target_p.insert_paragraph_before('1.3.2 Specific Objectives')
    h2_2.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else h2_2.style
    
    p1 = target_p.insert_paragraph_before('')
    p1.add_run('1. To analyze the current systems: ').bold = True
    p1.add_run('This involves conducting a comprehensive study of the existing manual and digital processes used for industrial attachment placement. It aims to identify pain points, inefficiencies, and the specific requirements of students, educational institutions, and companies to inform the new solution.')

    p2 = target_p.insert_paragraph_before('')
    p2.add_run('2. To design the system: ').bold = True
    p2.add_run('This entails developing the architectural framework, database schemas, and user interface (UI/UX) blueprints for the AI-powered matching platform. The design ensures scalable, secure, and intuitive interaction across different stakeholder portals while adhering to compliance and accessibility standards.')

    p3 = target_p.insert_paragraph_before('')
    p3.add_run('3. To implement the system: ').bold = True
    p3.add_run('This focuses on writing code to build the actual system, including the AI matching algorithm, automated document generation, M-Pesa mobile payment integrations, and intelligent chatbots. It also includes rigorous testing and deploying the final platform for stakeholder adoption.')

    doc.save(out_filepath)
    print("Successfully updated document.")
    return True

if __name__ == "__main__":
    update_objectives('/home/wakanda_forever/Desktop/AISHA/documentation/AISHA_Project_Proposal.docx', '/home/wakanda_forever/Desktop/AISHA/documentation/AISHA_Project_Proposal.docx')
